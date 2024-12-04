
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from django.contrib.auth.models import User

from table.models import PublicationType
from user.models import UserModel

from docx.shared import RGBColor

def set_bold(cell):
    """
    Устанавливает жирное начертание текста в ячейке.

    :param cell: Объект ячейки
    """
    # Получаем первый элемент run в ячейке
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True


def set_font_color(cell, rgb_color):
    """
    Устанавливает цвет текста в ячейке.

    :param cell: Объект ячейки
    :param rgb_color: Цвет в формате RGB, например, RGBColor(255, 0, 0) для красного
    """
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = rgb_color


from docx.shared import Inches, Pt

from utils.parseCoauthor import parse_json_to_dict


def set_vertical_alignment(cell, alignment="center"):
    """
    Устанавливает вертикальное выравнивание текста в ячейке.
    
    :param cell: Объект ячейки
    :param alignment: Значение выравнивания (top, center, bottom)
    """
    tc_pr = cell._element.get_or_add_tcPr()
    v_align = OxmlElement('w:vAlign')
    v_align.set(qn('w:val'), alignment)
    tc_pr.append(v_align)

from docx.oxml.ns import qn

from docx.enum.text import WD_ALIGN_PARAGRAPH

class DocxStyle:
    def set_vertical_alignment(self, cell, alignment="center"):
        """
        Устанавливает вертикальное выравнивание текста в ячейке.
        
        :param cell: Объект ячейки
        :param alignment: Значение выравнивания (top, center, bottom)
        """
        tc_pr = cell._element.get_or_add_tcPr()
        v_align = OxmlElement('w:vAlign')
        v_align.set(qn('w:val'), alignment)
        tc_pr.append(v_align)
    
    
    def set_horizontal_alignment(self, cell, alignment="center"):
        """
        Устанавливает горизонтальное выравнивание текста в ячейке.
        
        :param cell: Объект ячейки
        :param alignment: Значение выравнивания (left, center, right)
        """
        for paragraph in cell.paragraphs:
            if alignment == "center":
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif alignment == "right":
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    def set_font_color(self, cell, rgb_color):
        """
        Устанавливает цвет текста в ячейке.

        :param cell: Объект ячейки
        :param rgb_color: Цвет в формате RGB, например, RGBColor(255, 0, 0) для красного
        """
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = rgb_color
    
    
    def set_bold(self, cell):
        """
        Устанавливает жирное начертание текста в ячейке.

        :param cell: Объект ячейки
        """
        # Получаем первый элемент run в ячейке
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    
    def set_column_width(self, table, widths):
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = width
    
    
    def set_font(self, cell, font_name='Times New Roman', font_size=12):
        """Устанавливаем шрифт для текста в ячейке"""
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.name = font_name
                run.font.size = Pt(font_size)
                
                # Для совместимости с Word необходимо установить 'w:rFonts' через xml
                rFonts = run._element.rPr.rFonts
                rFonts.set(qn('w:eastAsia'), font_name)
    
    
    def set_landscape(self, section):
        """
        Устанавливает альбомную ориентацию для указанной секции документа Word.
        
        :param section: Объект секции документа Word
        """
        section_start = section._sectPr
        page_size = section_start.xpath('./w:pgSz')[0]
        page_size.set(qn('w:orient'), 'landscape')
        page_size.set(qn('w:w'), '16840')  # ширина страницы в EMU (21 см для альбомной ориентации)
        page_size.set(qn('w:h'), '11900')  # высота страницы в EMU (29.7 см для альбомной ориентации)

class DocxGenerator:
    def __init__(self, data, filename, user):
        self.data = data
        self.filename = filename
        self.docxStyle = DocxStyle()
        self.user = user
    


    def create_docx(self):
        doc = Document()
        
        
      
        self.docxStyle.set_landscape(doc.sections[0])
        
        # Даты
        dates = {rec["data"] for rec in self.data}
        # self.data.sort(key=lambda x: x["date"]) 
        
        # date_sorted = sorted(data, key=lambda x: x["data"])
 
        
        
        
        user_ids = {rec["for_user"] for rec in self.data}
        
        users = User.objects.filter(id__in=user_ids).values_list('id', 'username')
        
        # Создаем словарь для быстрого доступа к имени пользователя по его id
        user_dict = {user_id: username for user_id, username in users}

        # Заполняем список имен пользователей
        owner_name_recording = [user_dict[user_id] for user_id in user_ids]
        
       

        # Добавляем заголовок
     
        doc.add_heading(f"""
        В Национальном университете Узбекистана имени Мирзо Улугбека
        НИИ физики полупроводников и микроэлектроники
        { self.user.username or "Аноним" } , заведующий лабораторией «Возобновляемые источники энергии»
        СПИСОК НАУЧНЫХ ПУБЛИКАЦИЙ
        ({min(dates)} - {max(dates)} годы)
        """).runs[0].font.color.rgb = RGBColor(0, 0, 0) 
        
        
        # doc.add_heading("Записи пользователей:", level = 1).runs[0].font.color.rgb = RGBColor(0, 0, 0) 
        
        
        for i in range(len(owner_name_recording)):
            doc.add_paragraph(f"{i + 1}. {owner_name_recording[i]}")
 
        
        

        # Определяем заголовки таблицы
        headers = ['№', 'Название научной работы', 'Тип публикации', 'Информация об издании', 'Кол-во страниц', 'Соавторы']
        
        # Создаем таблицу с одной строкой для заголовков
        # table = doc.add_table(rows=len(self.data), cols=len(headers))
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        # table.font  = "Times New Roman" 
         
        

        
        section = doc.sections[0]
        page_width = section.page_width - section.left_margin - section.right_margin
        
        column_widths = [
            page_width * 0.01,  # № (1%)
            page_width * 0.1,   # Название научной работы (10%)
            page_width * 0.1,   # Тип публикации (10%)
            page_width * 0.35,  # Информация об издании (35%)
            page_width * 0.01,   # Кол-во страниц (1%)
            page_width * 0.39    # Соавторы (39%)
        ]

        


        section = doc.sections[-1]
        section.top_margin = Inches(0.8) #Верхний отступ
        section.bottom_margin = Inches(0.8) #Нижний отступ
        section.left_margin = Inches(0.3) #Отступ слева
        section.right_margin = Inches(0.3) #Отступ справа
                

        # Заполняем заголовки
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            self.docxStyle.set_bold(hdr_cells[i])  # Делаем заголовки жирными
            self.docxStyle.set_font_color(hdr_cells[i], RGBColor(0, 0, 0))
            self.docxStyle.set_horizontal_alignment(hdr_cells[i], 'center') 
            self.docxStyle.set_vertical_alignment(hdr_cells[i], 'center')
            self.docxStyle.set_font(hdr_cells[i])   
          
                   

        

        # Заполняем таблицу данными
        for index, item in enumerate(self.data):
            row_cells = table.add_row().cells
            row_cells[0].text = str(index + 1)
            row_cells[1].text = item.get('name', '')
            row_cells[2].text = str(PublicationType.objects.get(pk = int( item.get('Type', ''))))
            

           

            row_cells[3].text = f"{item.get('title', '')}, {str(item.get('data', ''))}, {str(item.get('tom', ''))}, {str(item.get('issue', ''))}"
            row_cells[4].text = f"от {str(item.get('page_start', ''))} до {str(item.get('page_end', ''))} всего {str(item.get('pages', ''))}"
          
            
            for i, key in enumerate(parse_json_to_dict(item.get("authors", []))):
                row_cells[5].text += f"{i + 1} Ф: {key['last_name']} И: {key['first_name']}; О: {key['patronymic']}\n\n" 
         
            for cell in row_cells:
                self.docxStyle.set_font(cell)
                self.docxStyle.set_horizontal_alignment(cell, 'center') 
                self.docxStyle.set_vertical_alignment(cell, 'center')
          
        self.docxStyle.set_column_width(table, column_widths)

        # Сохраняем документ
        doc.save(self.filename)

  