
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


from docx.shared import RGBColor

class DocxStyle:

    def set_style_global(self, doc, table):
        section = doc.sections[0]
        page_width = section.page_width - section.left_margin - section.right_margin
        

        # column_widths = [
        #     page_width * 0.01,  # № (1%)
        #     page_width * 0.1,   # Название научной работы (10%)
        #     page_width * 0.1,   # Тип публикации (10%)
        #     page_width * 0.39,  # Информация об издании (39%)
        #     page_width * 0.01,   # Кол-во страниц (1%)
        #     page_width * 0.39    # Соавторы (39%)
        # ]

      

        column_widths = [
            page_width * 0.01 ,  # № (1%)
            page_width ,   # Название научной работы (10%)
            page_width * 0.1,   # Тип публикации (10%)
            page_width,  # Информация об издании (39%)
            page_width * 0.2,   # Кол-во страниц (1%)
            page_width    # Соавторы (39%)
        ]

        section = doc.sections[-1]
        section.top_margin = Inches(0.8) #Верхний отступ
        section.bottom_margin = Inches(0.8) #Нижний отступ
        section.left_margin = Inches(0.3) #Отступ слева
        section.right_margin = Inches(0.3) #Отступ справа 

        self.set_column_width(table, column_widths)
        

    def set_style_title(self, text):
         
        # Настраиваем свойства текста
        run = text.runs[0]
        run.font.color.rgb = RGBColor(0, 0, 0)
        for run in text.runs:

            run.font.name = 'Times New Roman'
            run.font.size = Pt(12) 
            
            rFonts = run._element.rPr.rFonts
            rFonts.set(qn('w:eastAsia'), 'Times New Roman')

        # Настраиваем выравнивание заголовка
        text.alignment = WD_ALIGN_PARAGRAPH.CENTER
    


    
    def set_vertical_alignment(self, cell, alignment="center"):
        """
        Устанавливает вертикальное выравнивание текста в ячейке.
        
        :param cell: Объект ячейки
        :param alignment: Значение выравнивания (top, center, bottom)
        """
        tc_pr = cell._element.get_or_add_tcPr()
        v_align = OxmlElement('w:vAlign')
        v_align.set(qn('w:val'), alignment)
        tc_pr.append(v_align)
    
    
    def set_horizontal_alignment(self, cell, alignment="center"):
        """
        Устанавливает горизонтальное выравнивание текста в ячейке.
        
        :param cell: Объект ячейки
        :param alignment: Значение выравнивания (left, center, right)
        """
        for paragraph in cell.paragraphs:
            if alignment == "center":
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif alignment == "right":
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    def set_font_color(self, cell, rgb_color):
        """
        Устанавливает цвет текста в ячейке.

        :param cell: Объект ячейки
        :param rgb_color: Цвет в формате RGB, например, RGBColor(255, 0, 0) для красного
        """
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = rgb_color
    
    
    def set_bold(self, cell):
        """
        Устанавливает жирное начертание текста в ячейке.

        :param cell: Объект ячейки
        """
        # Получаем первый элемент run в ячейке
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    
    def set_column_width(self, table, widths):
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = width
    
    
    
    def set_font(self, cell, font_name='Times New Roman', font_size=12):
        """Устанавливаем шрифт для текста в ячейке"""
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.name = font_name
                run.font.size = Pt(font_size)
                
                # Для совместимости с Word необходимо установить 'w:rFonts' через xml
                rFonts = run._element.rPr.rFonts
                rFonts.set(qn('w:eastAsia'), font_name)
    
    
    def set_landscape(self, section):
        """
        Устанавливает альбомную ориентацию для указанной секции документа Word.
        
        :param section: Объект секции документа Word
        """
        section_start = section._sectPr
        page_size = section_start.xpath('./w:pgSz')[0]
        page_size.set(qn('w:orient'), 'landscape')
        page_size.set(qn('w:w'), '16840')  # ширина страницы в EMU (21 см для альбомной ориентации)
        page_size.set(qn('w:h'), '11900')  # высота страницы в EMU (29.7 см для альбомной ориентации)