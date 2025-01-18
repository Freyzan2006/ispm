


from docx import Document




from table.models import PublicationType





from utils.parseCoauthor import parse_json_to_dict

from download.utils.Generator.DocxStyle import DocxStyle
from download.utils.Generator.DocxContent import DocxContent

class DocxGenerator:
    def __init__(self, data, filename, user):
        self.doc = Document()
        self.docxStyle = DocxStyle()
        self.content = DocxContent(data, user)
        self.filename = filename
        # self.user = user
        # self.data = data
        
        
        
    

    def create_docx(self):
        self.doc = Document()
        
        
      
        self.docxStyle.set_landscape(self.doc.sections[0])
        
    
        # Добавляем заголовок
        heading = self.doc.add_heading(level=1)
        run = heading.add_run(self.content.get_title_text())
        
        # Устанавливаем шрифт "Times New Roman" и размер шрифта
        run.font.name = 'Times New Roman'
        # self.doc.add_page_break()


        self.docxStyle.set_style_title(text = heading)


        # self.content.get_heading_table()

        # self.docxStyle.set_style_heading_table()
      
        

 
        
        

        # Определяем заголовки таблицы
        headers = ['№', 'Название научной работы', 'Тип публикации', 'Информация об издании', 'Кол-во страниц', 'Соавторы']
        
   
        table = self.doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"


        

        
        # section = self.doc.sections[0]
        # page_width = section.page_width - section.left_margin - section.right_margin
        
        # column_widths = [
        #     page_width * 0.01,  # № (1%)
        #     page_width * 0.1,   # Название научной работы (10%)
        #     page_width * 0.1,   # Тип публикации (10%)
        #     page_width * 0.35,  # Информация об издании (35%)
        #     page_width * 0.01,   # Кол-во страниц (1%)
        #     page_width * 0.39    # Соавторы (39%)
        # ]

        

        # self.docxStyle.set_style_global(doc = self.doc)

        # section = self.doc.sections[-1]
        # section.top_margin = Inches(0.8) #Верхний отступ
        # section.bottom_margin = Inches(0.8) #Нижний отступ
        # section.left_margin = Inches(0.3) #Отступ слева
        # section.right_margin = Inches(0.3) #Отступ справа
        
        # gg
        self.content.create_heading_table(table = table, docxStyle = self.docxStyle, headers = headers)

        # Заполняем заголовки
        # hdr_cells = table.rows[0].cells
        # for i, header in enumerate(headers):
        #     hdr_cells[i].text = header
        #     self.docxStyle.set_bold(hdr_cells[i])  # Делаем заголовки жирными
        #     self.docxStyle.set_font_color(hdr_cells[i], RGBColor(0, 0, 0))
        #     self.docxStyle.set_horizontal_alignment(hdr_cells[i], 'center') 
        #     self.docxStyle.set_vertical_alignment(hdr_cells[i], 'center')
        #     self.docxStyle.set_font(hdr_cells[i])   
          
                   
        create_paginated_tables(doc = self.doc, table = table, data = self.content.data, rows_per_page = 3, docxStyle = self.docxStyle)
        

        # Заполняем таблицу данными
        # for index, item in enumerate(self.content.data):
        #     row_cells = table.add_row().cells
        #     row_cells[0].text = str(index + 1)
        #     row_cells[1].text = item.get('name', '')
        #     row_cells[2].text = str(PublicationType.objects.get(pk = int( item.get('Type', ''))))
            

           

        #     row_cells[3].text = f"{item.get('title', '')}, {str(item.get('data', ''))}, {str(item.get('tom', ''))}, {str(item.get('issue', ''))}"
        #     row_cells[4].text = f"от {str(item.get('page_start', ''))} до {str(item.get('page_end', ''))} всего {str(item.get('pages', ''))}"
          
            
        #     for i, key in enumerate(parse_json_to_dict(item.get("authors", []))):
        #         row_cells[5].text += f"{key['first_name'][0].upper()}. {key['patronymic'][0].upper()}. {key['last_name']}\n\n" 
         
        #     for index, cell in enumerate(row_cells):    
        #         self.docxStyle.set_font(cell)
        #         if index != 5:
        #             self.docxStyle.set_horizontal_alignment(cell, 'center') 
        #         self.docxStyle.set_vertical_alignment(cell, 'center')
        
       
        self.docxStyle.set_style_global(doc = self.doc, table = table)
        
       

        # Сохраняем документ
        self.doc.save(self.filename)


def create_paginated_tables(doc, table, data, rows_per_page, docxStyle):
    

    headers = ['№', 'Название научной работы', 'Тип публикации', 'Информация об издании', 'Кол-во страниц', 'Соавторы']
    chunks = [data[i:i + rows_per_page] for i in range(0, len(data), rows_per_page)]

    for chunk_index, chunk in enumerate(chunks):
        # Для второй и последующих страниц
        if chunk_index > 0:
            doc.add_page_break()
            
        
        # Добавляем заголовки только для первой страницы и новых таблиц
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            # docxStyle.set_bold(hdr_cells[i])
            docxStyle.set_horizontal_alignment(hdr_cells[i], 'center')
            docxStyle.set_vertical_alignment(hdr_cells[i], 'center')
            docxStyle.set_font(hdr_cells[i], font_name="Times New Roman", font_size=12)

        # Заполняем строки данными
        for index, item in enumerate(chunk):
            row_cells = table.add_row().cells
            row_cells[0].text = str(index + 1 + chunk_index * rows_per_page)  # Корректируем индексацию
            row_cells[1].text = item.get('name', '')
            publication_type_id = item.get('Type')
            if publication_type_id:
                try:
                    row_cells[2].text = str(PublicationType.objects.get(pk=int(publication_type_id)))
                except PublicationType.DoesNotExist:
                    row_cells[2].text = "Неизвестный тип"
            else:
                row_cells[2].text = "Не указано"

            row_cells[3].text = f"{item.get('title', '')}, {str(item.get('data', ''))}, {str(item.get('tom', ''))}, {str(item.get('issue', ''))}"
            # row_cells[4].text = f"от {str(item.get('page_start', ''))} до {str(item.get('page_end', ''))} всего {str(item.get('pages', ''))}"
            row_cells[4].text = f"{str(item.get('pages', ''))}"

            row_cells[5].text = ""
            for key in parse_json_to_dict(item.get("authors", [])):
                row_cells[5].text += f"{key['first_name'][0].upper()}. {key['patronymic'][0].upper()}. {key['last_name']}\n\n"

            for idx, cell in enumerate(row_cells):
                docxStyle.set_font(cell)
                if idx != 5:
                    docxStyle.set_horizontal_alignment(cell, 'center')
                docxStyle.set_vertical_alignment(cell, 'center')